import os
import subprocess
from docx import Document

def extract_text_from_word(file_path, libreoffice_appimage):
    ext = os.path.splitext(file_path)[1].lower()
    if ext not in [".doc", ".docx"]:
        raise ValueError("Only .doc and .docx supported")

    # Convert DOC → DOCX
    if ext == ".doc":
        outdir = os.path.dirname(file_path) or "."
        base = os.path.splitext(os.path.basename(file_path))[0]
        converted = os.path.join(outdir, base + ".docx")
        subprocess.run(
            [
                libreoffice_appimage,
                "--headless",
                "--convert-to",
                "docx",
                file_path,
                "--outdir",
                outdir
            ],
            check=True
        )
        file_path = converted

    doc = Document(file_path)
    result = []
    seen_lines = set()

    # Recursive table processor
    def process_table(table):
        for row in table.rows:
            cells = []
            for cell in row.cells:
                value = cell.text.strip()
                if value and value not in cells:
                    cells.append(value)
            if cells:
                line = " | ".join(cells)
                if line not in seen_lines:
                    result.append(line)
                    seen_lines.add(line)
            # Handle nested tables
            for nested_table in cell.tables:
                process_table(nested_table)

    # Process paragraphs and tables in order
    def process_elements(elements):
        for element in elements:
            tag = element.tag.split("}")[-1]
            if tag == "p":
                for p in doc.paragraphs:
                    if p._element == element:
                        text = p.text.strip()
                        if text and text not in seen_lines:
                            result.append(text)
                            seen_lines.add(text)
                        break
            elif tag == "tbl":
                for table in doc.tables:
                    if table._element == element:
                        process_table(table)
                        break

    # 1. Process main body
    process_elements(doc.element.body)

    # 2. Process headers and footers
    for section in doc.sections:
        process_elements(section.header._element)
        process_elements(section.footer._element)

    return "\n".join(result)

if __name__=="__main__":
    libreoffice_appimage = r"F:\GIT_CONTRIBUTION\PYTHONIFY\Advanced_Word_Extractor\LibreOffice-7.4.7.2.basic-x86_64.AppImage"
    file_path = r"D:\opt\OCD\Backups & Builds\ref_docs\New folder\AnmolBishnoi.docx"
    text = extract_text_from_word(file_path, libreoffice_appimage)
    print(text)